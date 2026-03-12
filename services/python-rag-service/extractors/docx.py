import logging
from docx import Document
from .base import TextExtractor

logger = logging.getLogger(__name__)


class DOCXExtractor(TextExtractor):
    """Extract text from DOCX files"""
    
    def extract(self, file_path: str) -> str:
        """Extract text from DOCX"""
        if not self.validate_file(file_path):
            raise ValueError(f"Invalid DOCX file: {file_path}")
        
        try:
            text = []
            doc = Document(file_path)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text)
            
            result = "\n".join(text)
            logger.info(f"Extracted {len(result)} characters from DOCX")
            return result
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to extract DOCX: {str(e)}")
